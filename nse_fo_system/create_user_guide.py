"""
NSE F&O Live Dashboard — User Guide Generator
Creates a professional Word document using python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
from docx.shared import Inches
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Color palette — RGBColor for fonts, hex strings for cell fills ────────────
DARK_BLUE   = RGBColor(0x1A, 0x3A, 0x5C)   # Header / Section titles
MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # Sub-headings
LIGHT_BLUE  = RGBColor(0xD5, 0xE8, 0xF4)   # Table header fill
GREEN       = RGBColor(0x00, 0x70, 0x00)   # Bullish
RED         = RGBColor(0xC0, 0x00, 0x00)   # Bearish
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)   # Caution / Warning
GRAY_BG     = RGBColor(0xF2, 0xF2, 0xF2)   # Alternate row
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY   = RGBColor(0x40, 0x40, 0x40)
GOLD        = RGBColor(0xFF, 0xD7, 0x00)

# Hex string versions for cell backgrounds (python-docx XML shading)
HEX_DARK_BLUE  = "1A3A5C"
HEX_MID_BLUE   = "2E75B6"
HEX_LIGHT_BLUE = "D5E8F4"
HEX_GREEN      = "007000"
HEX_RED        = "C00000"
HEX_ORANGE     = "C55A11"
HEX_GRAY_BG    = "F2F2F2"
HEX_WHITE      = "FFFFFF"
HEX_GOLD       = "FFD700"
HEX_DARK_GRAY  = "404040"


# ══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def rgb_to_hex(color) -> str:
    """Convert RGBColor to 6-char hex string without using .red/.green/.blue."""
    if isinstance(color, str):
        return color.lstrip('#')
    # RGBColor in python-docx stores its value as a 3-byte bytestring in ._val
    # but the safest cross-version approach is to use its __str__ if it returns hex,
    # otherwise fall back to parsing the XML representation.
    s = str(color)  # many versions return e.g. '1A3A5C'
    if len(s) == 6 and all(c in '0123456789ABCDEFabcdef' for c in s):
        return s.upper()
    # Fallback: use the theme_color or just return white
    return "FFFFFF"


def set_cell_bg(cell, color):
    """Set cell background colour. Accepts hex string like 'D5E8F4' or RGBColor."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = rgb_to_hex(color)
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size=4):
    """Add thin border to a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=None, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Write formatted text into a table cell."""
    cell.text = ''
    p   = cell.paragraphs[0]
    p.alignment = align
    # cell padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top','left','bottom','right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    '80')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return run


def add_heading(text, level=1, color=DARK_BLUE, size=None, space_before=18, space_after=6):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.font.color.rgb = color
    return p


def add_body(text, color=DARK_GRAY, size=10.5, space_after=4, bold=False, italic=False):
    """Add body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic
    return p


def add_bullet(text, indent=0, color=DARK_GRAY, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
        r1.font.color.rgb = color
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10)
        r2.font.color.rgb = color
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = color
    return p


def add_numbered(text, bold_prefix=None, color=DARK_GRAY):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(10)
        r1.font.color.rgb = color
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'; r2.font.size = Pt(10)
        r2.font.color.rgb = color
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'; run.font.size = Pt(10)
        run.font.color.rgb = color
    return p


def add_divider():
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_table(headers, rows, col_widths, header_bg=HEX_LIGHT_BLUE, alt_bg=HEX_GRAY_BG,
               header_color=DARK_BLUE, bold_header=True):
    """
    Create a styled table.
    headers  = list of header strings
    rows     = list of row tuples (each cell can be str or (str, RGBColor))
    col_widths = list of inches per column
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = Inches(w)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, "2E75B6", 6)
        cell_text(cell, hdr, bold=bold_header, color=header_color,
                  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = alt_bg if r_idx % 2 == 1 else HEX_WHITE
        tr = table.rows[r_idx + 1]
        for c_idx, cell_data in enumerate(row_data):
            cell = tr.cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC", 4)
            if isinstance(cell_data, tuple):
                txt, clr = cell_data
                cell_text(cell, txt, color=clr, size=10)
            else:
                cell_text(cell, cell_data, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def panel_heading(number, title, color=MID_BLUE):
    """Panel sub-heading with number badge."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"  Panel {number}:  ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = WHITE
    # shade the run background
    rPr = r1._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),  'clear')
    shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), rgb_to_hex(color))
    rPr.append(shd)
    r2 = p.add_run(f"  {title}")
    r2.bold = True
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = color
    return p


def scenario_box(number, title, color):
    """Scenario heading with colored left bar effect."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), rgb_to_hex(color))
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = p.add_run(f"SCENARIO {number}  —  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(12)
    r1.font.color.rgb = color
    r2 = p.add_run(title)
    r2.bold = True; r2.font.name = 'Calibri'; r2.font.size = Pt(12)
    r2.font.color.rgb = DARK_BLUE
    return p


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Title block with shaded background
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(48)
p_title.paragraph_format.space_after  = Pt(6)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:val'),  'clear')
shading_elm.set(qn('w:color'),'auto')
shading_elm.set(qn('w:fill'), '1A3A5C')
p_title._p.get_or_add_pPr().append(shading_elm)
r = p_title.add_run('NSE F&O LIVE DASHBOARD')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(28)
r.font.color.rgb = WHITE

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto')
shd2.set(qn('w:fill'),'2E75B6')
p_sub._p.get_or_add_pPr().append(shd2)
r2 = p_sub.add_run('Your Complete AI-Powered Options Trading Command Center')
r2.font.name = 'Calibri'; r2.font.size = Pt(14); r2.font.color.rgb = WHITE
r2.italic = True

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Tagline
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(6)
r3 = p3.add_run('Institutional-Grade Signals  |  Real-Time Data  |  Zero Manual Work')
r3.font.name = 'Calibri'; r3.font.size = Pt(11); r3.font.color.rgb = MID_BLUE
r3.bold = True

# Powered by
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(48)
r4 = p4.add_run('Powered by Zerodha Kite API  |  Built for NSE Options Traders')
r4.font.name = 'Calibri'; r4.font.size = Pt(10); r4.font.color.rgb = DARK_GRAY

add_divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PRODUCT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading('1.  What is NSE F&O Live Dashboard?', level=1)
add_divider()

add_body(
    'NSE F&O Live Dashboard is a real-time, browser-based trading intelligence platform '
    'built for Indian options traders. It connects directly to the Zerodha Kite API and '
    'delivers institutional-grade signals — the same signals used by hedge funds and '
    'professional traders — in a simple, clear interface.',
    size=10.5
)
add_body(
    'The dashboard auto-refreshes every 60 seconds without reloading the page. '
    'No manual calculations. No guesswork. Just clear BUY CE, BUY PE, SELL Iron Condor, '
    'or NO TRADE signals — backed by 7 independent data points working together.',
    size=10.5
)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_heading('Key Highlights', level=2, color=MID_BLUE, size=12, space_before=8)

make_table(
    headers=['Feature', 'Detail', 'Benefit'],
    col_widths=[1.8, 2.6, 2.6],
    rows=[
        ('Auto-Refresh',          'Every 60 seconds, no page reload',     'Always see live data — no F5 needed'),
        ('2 Tabs',                'Live Dashboard + Advanced Signals',     'Beginner and expert tools in one place'),
        ('7-Factor Signal Engine','PCR + OI + VIX + IV Rank + GEX + Max Pain + POC',
                                                                           'No single-indicator gambling'),
        ('Zerodha Integration',   'Direct Kite API — real NSE data',       'Accurate, real-time prices and OI'),
        ('Zero Manual Work',      'Fully automated calculations',          'Focus on trading, not spreadsheets'),
        ('Two Trading Styles',    'Intraday buying & Premium selling',     'Works for all types of traders'),
    ]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TAB 1: LIVE DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

add_heading('2.  Tab 1: Live Dashboard — 8 Panels Explained', level=1)
add_divider()
add_body(
    'Tab 1 is your primary trading screen. It refreshes every 60 seconds and shows '
    'everything you need to make an informed trade decision. All 8 panels work together — '
    'never rely on just one.',
    size=10.5
)

# Panel 1
panel_heading(1, 'Market Overview')
add_body('Shows live prices of NIFTY 50, NIFTY BANK, FIN NIFTY, and India VIX.')
make_table(
    headers=['VIX Level', 'Market Condition', 'What It Means', 'Action'],
    col_widths=[1.2, 1.8, 2.2, 1.8],
    rows=[
        (('Below 14',  GREEN),  'Very Calm',     'Fear is very low',              ('BUY options freely', GREEN)),
        (('14 – 20',   DARK_GRAY),'Normal',      'Standard market conditions',    ('Both buying & selling work', DARK_GRAY)),
        (('Above 20',  RED),    'High Fear',     'Panic in the market',           ('SELL premium — Iron Condor', ORANGE)),
    ]
)

# Panel 2
panel_heading(2, 'Trade Signal Engine  (Most Important Panel)')
add_body(
    'This is the heart of the dashboard. It automatically analyses all 7 factors and '
    'gives you one clear signal with entry price, target, stop loss, and confidence score.',
    bold=False
)
make_table(
    headers=['Signal', 'What It Means', 'What to Do'],
    col_widths=[1.6, 2.8, 2.6],
    rows=[
        (('BUY CE  🟢', GREEN),   'Market is bullish. Call options will go up.',
                                   'Buy the Call Option shown. Target +42%. SL -28%.'),
        (('BUY PE  🔴', RED),     'Market is bearish. Put options will go up.',
                                   'Buy the Put Option shown. Target +42%. SL -28%.'),
        (('SELL Iron Condor  💰', MID_BLUE),'Market is range-bound. Premium will decay.',
                                   'Sell CE + PE at strikes shown. Collect premium.'),
        (('NO TRADE  ⛔', DARK_GRAY),'Signals are mixed. Setup is unclear.',
                                   'Stay out. Protect capital. Wait for clarity.'),
    ]
)
add_body('Confidence Score Guide:', bold=True, color=DARK_BLUE)
add_bullet('Score below 30  →  NO TRADE. Not enough factors aligned.',    bold_prefix='')
add_bullet('Score 30 – 60   →  Moderate signal. Trade with smaller size.', bold_prefix='')
add_bullet('Score above 60  →  Strong signal. All major factors agree.',   bold_prefix='')

# Panel 3
panel_heading(3, 'OI Chain (Open Interest Chain)')
add_body('Shows Put and Call open interest at each strike. The ATM row is highlighted in green.')
add_body('Key signals in the BUILD column:', bold=True, color=DARK_BLUE)
make_table(
    headers=['BUILD Signal', 'Full Name', 'What It Means', 'Market Bias'],
    col_widths=[1.2, 1.8, 2.4, 1.6],
    rows=[
        (('FL', GREEN),  'Fresh Long',   'New buyers entering at this strike', ('BULLISH', GREEN)),
        (('FS', RED),    'Fresh Short',  'New sellers entering at this strike',('BEARISH', RED)),
        (('SC', GREEN),  'Short Cover',  'Previous sellers are buying back',   ('BULLISH', GREEN)),
        (('LU', ORANGE), 'Long Unwind',  'Previous buyers are exiting',        ('BEARISH', ORANGE)),
        (('Neutral', DARK_GRAY),'No Change','No significant OI movement',      ('WAIT', DARK_GRAY)),
    ]
)
add_bullet('Best setup: FL on CE at ATM + PCR rising = Strong BUY CE opportunity.')
add_bullet('Best setup: FS on PE at ATM + PCR falling = Strong BUY PE opportunity.')

# Panel 4
panel_heading(4, 'PCR — Put Call Ratio')
add_body('Shows PCR value for NIFTY and BANKNIFTY with a live trend arrow (▲ or ▼).')
make_table(
    headers=['PCR Value', 'Trend', 'Signal', 'Meaning'],
    col_widths=[1.2, 1.2, 1.8, 2.8],
    rows=[
        (('Above 1.2', GREEN),  '▲ Rising', ('BULLISH', GREEN),  'More puts being bought = strong support below'),
        (('0.8 – 1.2', DARK_GRAY),'→ Flat',('NEUTRAL', DARK_GRAY),'Balanced market. Wait for breakout.'),
        (('Below 0.8', RED),    '▼ Falling',('BEARISH', RED),    'More calls being bought = resistance above'),
    ]
)

# Panel 5
panel_heading(5, 'IV Rank & ATM Greeks')
add_body('IV Rank tells you whether options are cheap or expensive compared to the last 52 weeks.')
make_table(
    headers=['IV Rank', 'Options Are', 'Best Strategy'],
    col_widths=[1.6, 2.0, 3.4],
    rows=[
        (('Below 30%', GREEN),  'CHEAP — Low premium',  'BUY options. Good risk-reward ratio.'),
        (('30% – 60%', DARK_GRAY),'NORMAL', 'Follow other signals. Standard conditions.'),
        (('Above 60%', ORANGE), 'EXPENSIVE — High premium', 'SELL options. Iron Condor or Straddle.'),
    ]
)
add_body('Theta Clock: Shows how much rupee value options lose per day. For option sellers, this is your daily income from time decay.')

# Panel 6
panel_heading(6, 'GEX — Gamma Exposure (Institutional Signal)')
add_body(
    'GEX shows the net hedging obligation of market makers. It tells you whether the market '
    'will stay in a range or break out with a big move.'
)
make_table(
    headers=['GEX Reading', 'Regime', 'What Happens', 'Best Strategy'],
    col_widths=[1.4, 1.6, 2.2, 2.0],
    rows=[
        (('Positive (+)', GREEN), 'RANGE BOUND',       'MMs stabilize price. Range trading.',     ('Iron Condor / Sell Strangle', GREEN)),
        (('Near Zero',DARK_GRAY), 'TRANSITION',        'Regime about to change. Be careful.',     ('Reduce position size', ORANGE)),
        (('Negative (−)', RED),   'VOLATILE/TRENDING', 'MMs amplify moves. Big move expected.',   ('Buy CE or PE directionally', RED)),
    ]
)
add_bullet('Gamma Wall: The strike NIFTY is most magnetically attracted to. Strong support/resistance.')
add_bullet('Flip Level: If price crosses below this level, the market turns from range-bound to volatile.')

# Panel 7
panel_heading(7, 'Volume Profile — POC · VAH · VAL')
add_body(
    'Shows a horizontal bar chart of where the most trading volume occurred. '
    'Three key levels are calculated: POC, VAH, and VAL. '
    'Use the session selector to switch between Today / Weekly / Monthly view.'
)
make_table(
    headers=['Level', 'Full Name', 'What It Is', 'How to Trade It'],
    col_widths=[0.8, 1.8, 2.2, 2.2],
    rows=[
        (('POC', ORANGE),  'Point of Control',  'Price where MOST volume traded',     'Price above POC = bullish. Below = bearish.'),
        (('VAH', MID_BLUE),'Value Area High',   'Top of 70% volume zone',             'Price breaking above VAH = strong breakout → BUY CE'),
        (('VAL', MID_BLUE),'Value Area Low',    'Bottom of 70% volume zone',          'Price breaking below VAL = breakdown → BUY PE'),
    ]
)
add_body('Quick Rules:', bold=True, color=DARK_BLUE)
add_bullet('Price above POC  →  Bullish. POC = support below you. Buy CE on dips.', bold_prefix='')
add_bullet('Price below POC  →  Bearish. POC = resistance above you. Buy PE on bounces.', bold_prefix='')
add_bullet('Price inside Value Area  →  Range-bound. Sell Iron Condor.', bold_prefix='')
add_bullet('Price breaks outside Value Area  →  Strong directional move starting. Trade the breakout.', bold_prefix='')

# Panel 8
panel_heading(8, 'OI Buildup Analysis')
add_body(
    'Filters the OI chain to show only the strikes where significant changes are happening. '
    'Helps you see where big money is entering or exiting.'
)
add_bullet('Fresh Long at ATM + PCR rising  →  Strongest BUY CE confirmation.')
add_bullet('Fresh Short at ATM + PCR falling  →  Strongest BUY PE confirmation.')
add_bullet('Short Cover appearing  →  Trend may be reversing upward. Reduce short positions.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TAB 2: ADVANCED SIGNALS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('3.  Tab 2: Advanced Signals — 5 Institutional Tools', level=1)
add_divider()

p_intro = doc.add_paragraph()
p_intro.paragraph_format.space_after = Pt(10)
shd_intro = OxmlElement('w:shd')
shd_intro.set(qn('w:val'),'clear'); shd_intro.set(qn('w:color'),'auto')
shd_intro.set(qn('w:fill'),'EBF3FB')
p_intro._p.get_or_add_pPr().append(shd_intro)
r_intro = p_intro.add_run(
    '  These 5 signals are used by hedge funds and institutional traders. '
    'They are now available to retail traders through this dashboard.  '
)
r_intro.font.name = 'Calibri'; r_intro.font.size = Pt(10.5)
r_intro.font.color.rgb = DARK_BLUE; r_intro.italic = True

# Panel 1 — SMI
panel_heading(1, 'Smart Money Index (SMI)', color=MID_BLUE)
add_body(
    'Tracks the difference between retail trading behaviour (morning 9:15–9:45) and '
    'institutional behaviour (closing 3:00–3:30). Retail traders panic in the morning. '
    'Institutions quietly position themselves in the last 30 minutes.'
)
add_body('Formula:  SMI = Previous SMI  –  Morning Move  +  Evening Move', bold=True, color=MID_BLUE)
doc.add_paragraph().paragraph_format.space_after = Pt(3)
make_table(
    headers=['Morning Move', 'Evening Move', 'SMI Signal', 'Action Tomorrow'],
    col_widths=[1.5, 1.5, 2.2, 2.0],
    rows=[
        (('DOWN ▼', RED),    ('UP ▲', GREEN),   ('INSTITUTIONS BUYING', GREEN),        ('Buy CE tomorrow morning', GREEN)),
        (('UP ▲', GREEN),    ('DOWN ▼', RED),   ('DISTRIBUTION', RED),                 ('Buy PE tomorrow morning', RED)),
        (('UP ▲', GREEN),    ('UP ▲', GREEN),   ('BULLISH MOMENTUM', GREEN),           ('Stay / Add Long', GREEN)),
        (('DOWN ▼', RED),    ('DOWN ▼', RED),   ('BEARISH PRESSURE', RED),             ('Stay / Add Short', RED)),
    ]
)
add_body(
    'Golden Rule: SMI rising while NIFTY price is falling = Institutions silently buying = '
    'Tomorrow will be UP. This is one of the most reliable next-day signals available.',
    italic=True, color=MID_BLUE
)

# Panel 2 — Gamma Acceleration
panel_heading(2, 'Gamma Acceleration  (dGEX/dt)', color=MID_BLUE)
add_body(
    'Measures the SPEED at which GEX is changing every minute. GEX alone tells you '
    'the current regime. Gamma Acceleration tells you when the regime is ABOUT TO CHANGE '
    '— before it actually happens.'
)
add_body('Think of it as: GEX = weather today. Gamma Acceleration = weather forecast for next hour.', italic=True, color=DARK_GRAY)
make_table(
    headers=['What You See', 'What It Means', 'Immediate Action'],
    col_widths=[2.0, 2.5, 2.5],
    rows=[
        (('Flip ETA < 15 min', RED),    'URGENT: Regime flip imminent',         ('CLOSE all short option positions NOW', RED)),
        (('GEX Decaying fast', ORANGE), 'Market about to turn volatile',        ('Stop selling. Prepare to buy CE or PE', ORANGE)),
        (('GEX Stable Positive',GREEN), 'Range-bound confirmed',                ('Continue Iron Condor. Stay in trade.', GREEN)),
        (('GEX Recovering', MID_BLUE),  'Volatility calming down',              ('Start building selling positions', MID_BLUE)),
    ]
)

# Panel 3 — Strike Pinning
panel_heading(3, 'Strike Pinning Probability', color=MID_BLUE)
add_body(
    'Shows the probability that NIFTY will close AT each strike on expiry day. '
    'Market makers hold massive hedges at specific strikes. Their re-hedging creates '
    'a "gravity pull" that draws NIFTY price toward those strikes near expiry.'
)
add_body('Formula:  Pin Probability = |Net GEX at Strike| / Total |Net GEX| × 100', bold=True, color=MID_BLUE)
add_body('Example Trade:', bold=True, color=DARK_BLUE)
add_bullet('Strike 24000 shows 55% pin probability.')
add_bullet('Action: SELL 24000 CE + SELL 24000 PE (Straddle).')
add_bullet('If NIFTY closes near 24000 at expiry, both options expire near zero → full premium collected.')
add_bullet('Stop Loss: If NIFTY moves more than 100 points away from 24000, exit immediately.')
add_body('Best Entry Window: Wednesday 2:00 PM to Thursday 10:30 AM (for weekly expiry).', italic=True, color=DARK_GRAY)

# Panel 4 — Expected Move
panel_heading(4, 'Expected Move Calculator', color=MID_BLUE)
add_body(
    'Automatically calculates how much NIFTY is expected to move by expiry, '
    'using the ATM straddle price. The dashboard shows the Upper Limit, Lower Limit, '
    'and an automatic Iron Condor suggestion.'
)
add_body('Formula:  Expected Move = ATM Call Premium + ATM Put Premium', bold=True, color=MID_BLUE)
add_body('85% probability NIFTY stays within the Upper and Lower limits shown.', italic=True, color=DARK_GRAY)
make_table(
    headers=['Straddle Size', 'Market Expectation', 'Best Strategy'],
    col_widths=[1.8, 2.4, 3.0],
    rows=[
        (('Above 2.5% of spot', RED),   'Big move coming (High IV)',   ('Buy options directionally (CE or PE)', DARK_GRAY)),
        (('1% – 2.5%', DARK_GRAY),      'Normal conditions',           ('Both buying and selling viable', DARK_GRAY)),
        (('Below 1%', GREEN),           'Calm, low volatility',        ('Sell Iron Condor — highest win rate', GREEN)),
    ]
)

# Panel 5 — Cross-Asset
panel_heading(5, 'Cross-Asset Signals', color=MID_BLUE)
add_body(
    'NIFTY does not move alone. These correlated markets give early signals '
    '— often 15–30 minutes before NIFTY reacts.'
)
make_table(
    headers=['Signal', 'What It Indicates', 'NIFTY Impact'],
    col_widths=[2.0, 2.8, 2.2],
    rows=[
        ('India VIX falling below 14', 'Market fear reducing',                   ('POSITIVE — Buy strategies safe', GREEN)),
        ('BankNifty outperforming',    'FII buying financials',                   ('POSITIVE — Broad rally likely', GREEN)),
        ('Rupee strengthening',        'FII buying Indian equities',              ('POSITIVE — NIFTY support', GREEN)),
        ('Rupee weakening',            'FII selling Indian equities',             ('NEGATIVE — NIFTY under pressure', RED)),
        ('VIX above 20',               'High fear — options overpriced',          ('Sell premium, avoid buying', ORANGE)),
    ]
)
add_bullet('3+ signals BULLISH  →  Strong long setup. Buy CE with higher confidence.')
add_bullet('3+ signals BEARISH  →  Strong short setup. Buy PE with higher confidence.')
add_bullet('Check SGX Nifty before 9 AM on sgxnifty.com — gives 15-min advance warning of gap-up / gap-down open.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TRADING STRATEGY GUIDE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('4.  When to BUY, SELL, and Stay Out', level=1)
add_divider()
add_body(
    'Use the checklist below before entering any trade. The more conditions that are '
    'met, the higher your probability of success. Never enter a trade if fewer than '
    '4 conditions are met.',
    size=10.5
)

# Scenario 1
scenario_box(1, 'BUY CE  (Bullish Intraday Trade)', GREEN)
add_body('Enter when ALL of these are true:', bold=True, color=DARK_GRAY)
add_numbered('Trade Signal panel shows  BUY CE', bold_prefix='Trade Signal: ')
add_numbered('PCR is above 1.2  AND  trend arrow is rising (▲)', bold_prefix='PCR: ')
add_numbered('GEX Regime is  VOLATILE / TRENDING  (negative or near-zero)', bold_prefix='GEX: ')
add_numbered('OI Build at ATM shows  Fresh Long (FL)', bold_prefix='OI Build: ')
add_numbered('India VIX is below 18', bold_prefix='VIX: ')
add_numbered('NIFTY spot is  ABOVE  the POC level (Volume Profile)', bold_prefix='Volume Profile: ')
add_numbered('SMI is rising (Tab 2)', bold_prefix='SMI (Tab 2): ')
make_table(
    headers=['Entry', 'Target', 'Stop Loss', 'Timeframe', 'Exit Rule'],
    col_widths=[1.4, 1.3, 1.3, 1.5, 1.5],
    rows=[(
        'As shown in Trade Signal panel',
        ('+42% from entry', GREEN),
        ('-28% from entry', RED),
        'Intraday only',
        'Exit before 3:20 PM',
    )]
)

# Scenario 2
scenario_box(2, 'BUY PE  (Bearish Intraday Trade)', RED)
add_body('Enter when ALL of these are true:', bold=True, color=DARK_GRAY)
add_numbered('Trade Signal panel shows  BUY PE', bold_prefix='Trade Signal: ')
add_numbered('PCR is below 0.8  AND  trend arrow is falling (▼)', bold_prefix='PCR: ')
add_numbered('GEX Regime is  VOLATILE / TRENDING', bold_prefix='GEX: ')
add_numbered('OI Build at ATM shows  Fresh Short (FS)', bold_prefix='OI Build: ')
add_numbered('India VIX is above 15 (some fear confirms bearish direction)', bold_prefix='VIX: ')
add_numbered('NIFTY spot is  BELOW  the POC level (Volume Profile)', bold_prefix='Volume Profile: ')
add_numbered('SMI is falling (Tab 2)', bold_prefix='SMI (Tab 2): ')
make_table(
    headers=['Entry', 'Target', 'Stop Loss', 'Timeframe', 'Exit Rule'],
    col_widths=[1.4, 1.3, 1.3, 1.5, 1.5],
    rows=[(
        'As shown in Trade Signal panel',
        ('+42% from entry', GREEN),
        ('-28% from entry', RED),
        'Intraday only',
        'Exit before 3:20 PM',
    )]
)

# Scenario 3
scenario_box(3, 'SELL Iron Condor  (Range-Bound Premium Collection)', MID_BLUE)
add_body('Enter when ALL of these are true:', bold=True, color=DARK_GRAY)
add_numbered('Trade Signal panel shows  SELL — Iron Condor', bold_prefix='Trade Signal: ')
add_numbered('IV Rank is above 60% (options are expensive — best time to sell)', bold_prefix='IV Rank: ')
add_numbered('GEX Regime is  RANGE BOUND  (positive GEX)', bold_prefix='GEX: ')
add_numbered('India VIX is between 14 – 20', bold_prefix='VIX: ')
add_numbered('NIFTY is inside the Value Area on Volume Profile', bold_prefix='Volume Profile: ')
add_numbered('Expected Move straddle is below 1.5% of spot (Tab 2)', bold_prefix='Expected Move: ')
make_table(
    headers=['Sell CE', 'Sell PE', 'Collect Premium', 'Stop Loss Rule', 'Timeframe'],
    col_widths=[1.4, 1.4, 1.4, 2.2, 1.2],
    rows=[(
        'At strike shown',
        'At strike shown',
        ('Full premium = Max profit', GREEN),
        ('Exit if either leg 2× original price', RED),
        'Weekly / Swing',
    )]
)
add_body('You profit when NIFTY stays between the two sold strikes at expiry.', italic=True, color=DARK_GRAY)

# Scenario 4
scenario_box(4, 'NO TRADE  —  Stay Out and Protect Capital', DARK_GRAY)
add_body('Do NOT trade when any of these conditions exist:', bold=True, color=DARK_GRAY)
add_bullet('Trade Signal shows NO TRADE or confidence score is below 30.')
add_bullet('Signals are mixed — some bullish and some bearish. Setup is unclear.')
add_bullet('VIX spikes suddenly above 25 (panic mode — option spreads widen, slippage increases).')
add_bullet('It is expiry morning before 10:30 AM (option prices are highly erratic).')
add_bullet('Budget, RBI Policy, Election results, or major global event is scheduled today.')
add_bullet('Gamma Acceleration shows Flip ETA under 15 minutes — regime about to change.')

p_rule = doc.add_paragraph()
p_rule.paragraph_format.space_before = Pt(8)
p_rule.paragraph_format.space_after  = Pt(8)
shd_rule = OxmlElement('w:shd')
shd_rule.set(qn('w:val'),'clear'); shd_rule.set(qn('w:color'),'auto')
shd_rule.set(qn('w:fill'),'FFF2CC')
p_rule._p.get_or_add_pPr().append(shd_rule)
r_rule = p_rule.add_run(
    '  Golden Rule:  Protecting capital is more important than catching every trade. '
    'Missing a trade costs nothing. A bad trade costs real money. When in doubt — STAY OUT.  '
)
r_rule.bold = True; r_rule.font.name = 'Calibri'; r_rule.font.size = Pt(11)
r_rule.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('5.  Quick Reference Cheat Sheet', level=1)
add_divider()

add_heading('Signal Colours and Meanings', level=2, color=MID_BLUE, size=12, space_before=8)
make_table(
    headers=['Colour', 'Meaning', 'Immediate Action'],
    col_widths=[1.4, 2.6, 3.0],
    rows=[
        (('Green  ✅', GREEN),   'Bullish signal — conditions favour an upward move', 'Favour BUY CE or stay long'),
        (('Red  ❌', RED),       'Bearish signal — conditions favour a downward move','Favour BUY PE or stay short'),
        (('Yellow  ⚠', ORANGE), 'Caution / Neutral — signals are mixed',             'Wait for confirmation. Do not force a trade.'),
        (('Gray  ⚪', DARK_GRAY),'No data or signal inactive',                        'Ignore this indicator for now'),
    ]
)

add_heading('Key Levels to Watch Every Day', level=2, color=MID_BLUE, size=12, space_before=8)
make_table(
    headers=['Level', 'What It Is', 'Why It Matters for Trading'],
    col_widths=[1.6, 2.2, 3.2],
    rows=[
        ('POC',                  'Point of Control (Volume Profile)', 'Strongest support/resistance. Price returns here.'),
        ('VAH',                  'Value Area High',                   'Break above = strong bullish breakout. Buy CE.'),
        ('VAL',                  'Value Area Low',                    'Break below = strong bearish breakdown. Buy PE.'),
        ('Gamma Wall',           'Highest GEX strike',                'NIFTY is magnetically attracted here. S/R zone.'),
        ('Flip Level',           'GEX zero crossing strike',          'Below this = Volatile. Above this = Range-bound.'),
        ('Max Pain',             'Lowest loss strike for sellers',    'NIFTY tends to close near this on expiry day.'),
        ('Expected Move Upper',  'Spot + ATM Straddle',               'Sell CE above this level. High win rate.'),
        ('Expected Move Lower',  'Spot – ATM Straddle',               'Sell PE below this level. High win rate.'),
    ]
)

add_heading('VIX-Based Strategy Guide', level=2, color=MID_BLUE, size=12, space_before=8)
make_table(
    headers=['VIX Level', 'Market Condition', 'Best Strategy', 'Avoid'],
    col_widths=[1.3, 1.8, 2.7, 1.7],
    rows=[
        (('Below 14', GREEN),   'Very calm',       'Buy options — they are cheap. Directional trades.',     'Selling premium (low reward)'),
        (('14 – 20', DARK_GRAY),'Normal',           'Both buying and selling work. Follow other signals.',  'Overtrading'),
        (('20 – 25', ORANGE),   'Elevated fear',    'Sell premium. Iron Condor. Straddle sell.',            'Buying options (expensive)'),
        (('Above 25', RED),     'Panic / Crisis',   'No new trades. Protect open positions.',               'Any new trade entry'),
    ]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading('6.  Key Terms Explained (Glossary)', level=1)
add_divider()

terms = [
    ('CE (Call Option)',       'Right to buy NIFTY at a specific price. Profit when NIFTY goes UP.'),
    ('PE (Put Option)',        'Right to sell NIFTY at a specific price. Profit when NIFTY goes DOWN.'),
    ('ATM (At The Money)',     'The strike price closest to the current NIFTY price.'),
    ('OTM (Out of The Money)', 'Strike above spot for CE, or below spot for PE. Cheaper but expires worthless more often.'),
    ('OI (Open Interest)',     'Total number of active option contracts. High OI = more institutional interest at that level.'),
    ('PCR (Put Call Ratio)',   'PE OI divided by CE OI. Above 1.2 = Bullish. Below 0.8 = Bearish.'),
    ('IV (Implied Volatility)','The market\'s expectation of future price movement. High IV = Options are expensive.'),
    ('IV Rank',                'Where current IV stands in its 52-week range. 0% = Historically cheapest. 100% = Historically most expensive.'),
    ('GEX (Gamma Exposure)',   'Net hedging obligation of market makers. Positive = Stabilising force. Negative = Amplifying force.'),
    ('POC (Point of Control)', 'The price level where the most volume traded in a session. Acts as the strongest support/resistance.'),
    ('VAH (Value Area High)',  'Upper boundary of the price zone where 70% of trading activity occurred.'),
    ('VAL (Value Area Low)',   'Lower boundary of the same 70% value area. Price inside VA = range-bound behaviour.'),
    ('Iron Condor',            'A strategy where you sell one CE and one PE at different strikes. Maximum profit if NIFTY stays between both strikes at expiry.'),
    ('Straddle (Sell)',        'Selling one ATM CE and one ATM PE at the same strike. Maximum profit if NIFTY closes exactly at that strike on expiry.'),
    ('SMI (Smart Money Index)','Tracks institutional vs retail activity. Rising SMI = institutions are quietly accumulating.'),
    ('Gamma Acceleration',     'The speed at which GEX is changing per minute. Predicts a regime change (range to volatile) before it happens.'),
    ('Pin Probability',        'The probability that NIFTY will close AT a specific strike on expiry day, based on GEX distribution.'),
    ('Expected Move',          'ATM Call + ATM Put premium = how much the market expects NIFTY to move by expiry. 85% confidence interval.'),
    ('Theta Decay',            'The daily reduction in option premium due to time passing. Option sellers benefit from this every day.'),
    ('Flip Level',             'The GEX strike where market behaviour flips from range-bound to volatile. Critical level to monitor.'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{term}:  ")
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(10)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(definition)
    r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DISCLAIMER
# ══════════════════════════════════════════════════════════════════════════════

add_heading('7.  Important Disclaimer', level=1, color=RED)
add_divider()

p_disc = doc.add_paragraph()
p_disc.paragraph_format.space_after = Pt(8)
shd_d = OxmlElement('w:shd')
shd_d.set(qn('w:val'),'clear'); shd_d.set(qn('w:color'),'auto')
shd_d.set(qn('w:fill'),'FFF0F0')
p_disc._p.get_or_add_pPr().append(shd_d)
r_d = p_disc.add_run(
    '  This dashboard is an analytical and educational tool designed to assist in '
    'trading decision-making. It does NOT guarantee profits. Options trading involves '
    'significant financial risk and is not suitable for all investors. Past signal '
    'accuracy does not guarantee future performance. Always use proper position sizing '
    'and strict stop losses. Never trade with money you cannot afford to lose. '
    'Consult a SEBI-registered financial advisor before making any investment decisions.  '
)
r_d.font.name = 'Calibri'; r_d.font.size = Pt(10.5); r_d.font.color.rgb = RED

add_body('Key Reminders:', bold=True, color=DARK_BLUE)
add_bullet('Always place a stop loss order immediately after entering a trade — no exceptions.')
add_bullet('Never risk more than 1–2% of your total trading capital on a single trade.')
add_bullet('Do not average down on losing trades.')
add_bullet('The dashboard provides signals, not certainty. Markets can and do behave unexpectedly.')
add_bullet('Past performance of any signal is not a guarantee of future results.')


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER — all pages
# ══════════════════════════════════════════════════════════════════════════════
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.paragraph_format.space_before = Pt(0)
footer_para.paragraph_format.space_after  = Pt(0)

# Thin top border on footer
pPr_f = footer_para._p.get_or_add_pPr()
pBdr_f = OxmlElement('w:pBdr')
top_f = OxmlElement('w:top')
top_f.set(qn('w:val'),   'single')
top_f.set(qn('w:sz'),    '4')
top_f.set(qn('w:space'), '1')
top_f.set(qn('w:color'), '2E75B6')
pBdr_f.append(top_f)
pPr_f.append(pBdr_f)

from docx.oxml import OxmlElement
fld_char1 = OxmlElement('w:fldChar')
fld_char1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fld_char2 = OxmlElement('w:fldChar')
fld_char2.set(qn('w:fldCharType'), 'end')

run_f1 = footer_para.add_run('NSE F&O Live Dashboard  |  Powered by Zerodha Kite API  |  For Educational Use  |  Page ')
run_f1.font.name = 'Calibri'; run_f1.font.size = Pt(8.5)
run_f1.font.color.rgb = DARK_GRAY

run_f2 = footer_para.add_run()
run_f2.font.name = 'Calibri'; run_f2.font.size = Pt(8.5)
run_f2.font.color.rgb = DARK_GRAY
run_f2._r.append(fld_char1)
run_f2._r.append(instrText)
run_f2._r.append(fld_char2)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r'D:\HDFC\nse_fo_system\NSE_FO_Dashboard_User_Guide.docx'
doc.save(output_path)
print(f"Document saved: {output_path}")
