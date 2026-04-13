"""
NSE F&O Dashboard — Word Document Generator
Ek click mein poora presentation document ban jayega
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Color palette ─────────────────────────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x0D, 0x1B, 0x2A)   # dark navy
C_BLUE        = RGBColor(0x00, 0x5C, 0xA8)   # professional blue
C_ACCENT      = RGBColor(0x00, 0xA8, 0xE8)   # cyan accent
C_GREEN       = RGBColor(0x00, 0x96, 0x60)   # green
C_RED         = RGBColor(0xCC, 0x00, 0x00)   # red
C_GOLD        = RGBColor(0xD4, 0xA0, 0x00)   # gold
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GRAY  = RGBColor(0xF2, 0xF4, 0xF7)
C_MID_GRAY    = RGBColor(0x55, 0x65, 0x77)
C_TEXT        = RGBColor(0x1A, 0x1A, 0x2E)
C_HEADING_BG  = RGBColor(0x00, 0x5C, 0xA8)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(11.69)  # A4 landscape
section.page_height = Inches(8.27)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   val.get('val',   'single'))
            border.set(qn('w:sz'),    val.get('sz',    '6'))
            border.set(qn('w:space'), val.get('space', '0'))
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def heading_para(text, size=26, color=C_WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.bold   = bold
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    return p

def body_para(text, size=11, color=C_TEXT, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    return p

def bullet_para(text, size=11, color=C_TEXT, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = C_BLUE
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.color.rgb = color
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p

def section_divider(title, bg_hex="005CA8", text_color=C_WHITE):
    """Full-width colored section header."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(9.5)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"  {title}")
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = text_color
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def two_col_table(rows_data, col1_width=2.5, col2_width=7.0, header=None):
    """Simple 2-column table."""
    num_cols = 2
    tbl = doc.add_table(rows=0, cols=num_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        row = tbl.add_row()
        for i, h in enumerate(header):
            c = row.cells[i]
            set_cell_bg(c, "005CA8")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = C_WHITE

    for i, (k, v) in enumerate(rows_data):
        row = tbl.add_row()
        bg = "F2F4F7" if i % 2 == 0 else "FFFFFF"
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, "EAF0F8")

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = C_BLUE

        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(10)
        r1.font.color.rgb = C_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def info_box(lines, bg_hex="EAF4FF", border_hex="005CA8"):
    """Shaded info box."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = C_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 — COVER
# ══════════════════════════════════════════════════════════════════════════════

# Big title block
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cover_tbl.rows[0].cells[0]
set_cell_bg(cell, "0D1B2A")

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("NSE F&O Live Trading Dashboard")
run.font.bold = True
run.font.size = Pt(28)
run.font.color.rgb = C_ACCENT

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run("Institutional-Grade Tools • Real-Time Signals • Smart Strike Selection")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
r3 = p3.add_run(f"Version 2.1  |  Zerodha Kite API  |  Streamlit Browser UI  |  {datetime.date.today().strftime('%B %Y')}")
r3.font.size = Pt(10)
r3.font.color.rgb = C_MID_GRAY

doc.add_paragraph()

# 3-column intro summary
summary = doc.add_table(rows=1, cols=3)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER

cols_data = [
    ("005CA8", "📊", "What It Is",
     "A fully automated NSE F&O analysis dashboard that reads live Zerodha data and generates actionable BUY / SELL / NO TRADE signals with specific strike prices, entry, target, and stop-loss."),
    ("006644", "🎯", "Who It's For",
     "Active options traders who want institutional-grade signals without paying for premium services. Works for NIFTY, BANKNIFTY, and FINNIFTY on weekly & monthly expiry."),
    ("7B2D00", "🔑", "Key Edge",
     "Combines GEX (Gamma Exposure), PCR, OI Build, IV Rank, VIX, and Max Pain into a single confidence score — this combination is NOT available in Sensibull, Opstra, or any free tool."),
]

for i, (bg, icon, title, text) in enumerate(cols_data):
    c = summary.rows[0].cells[i]
    set_cell_bg(c, bg)
    ph = c.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.space_before = Pt(8)
    rh = ph.add_run(f"{icon}  {title}")
    rh.font.bold = True
    rh.font.size = Pt(12)
    rh.font.color.rgb = C_WHITE

    pt = c.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pt.paragraph_format.space_after = Pt(8)
    rt = pt.add_run(f"\n{text}")
    rt.font.size = Pt(9)
    rt.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — WHAT WAS BUILT
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 1 — What Was Built  (Problem → Solution)")

body_para(
    "The dashboard was built from scratch to solve 8 real problems that traders face with existing tools. "
    "Each problem was identified, root-caused, and fixed with a specific technical solution.",
    size=11, color=C_TEXT
)

problems = [
    ("Dashboard loading 15–25 seconds",
     "ROOT CAUSE: kite.instruments('NFO') was being called 5–6 times every refresh, each call = 2–3 sec\n"
     "FIX: 30-minute instruments cache + parallel ThreadPoolExecutor(7 workers)\n"
     "RESULT: Load time reduced to 3–5 seconds (5x faster)"),
    ("PCR values wrong (20.77, 16.50)",
     "ROOT CAUSE: Deep OTM strikes have tiny CE OI → PE/CE ratio astronomically high\n"
     "FIX: Calculate PCR only from ATM ±30 strikes (1 API batch call)\n"
     "RESULT: Accurate PCR aligned with market reality"),
    ("OI Change always +0 after restart",
     "ROOT CAUSE: prev_oi dict was in-memory only, wiped on every restart\n"
     "FIX: Persist OI snapshot to data/prev_oi.pkl — survives restarts\n"
     "RESULT: Real OI change visible even on first data pull after restart"),
    ("ATM row invisible in OI chain",
     "ROOT CAUSE: Green highlight + dark text = unreadable. Strike label cut off\n"
     "FIX: Explicit white text (#ffffff) in highlight CSS, removed arrow symbols\n"
     "RESULT: ATM row clearly highlighted and readable"),
    ("UOA showing 'Scanning...' forever",
     "ROOT CAUSE: First baseline scan takes 60 seconds + 11 symbols too many\n"
     "FIX: Reduced to 2 symbols (NIFTY, BANKNIFTY), threshold from 5x → 2x\n"
     "RESULT: UOA results appear within 60 seconds"),
    ("start_web.bat window closing instantly",
     "ROOT CAUSE: Python/streamlit not in PATH, encoding error (ΓÇö in CMD)\n"
     "FIX: chcp 65001, proper error messages with pause, Y/N token prompt\n"
     "RESULT: BAT file guides user step by step, never closes silently"),
    ("No Trade Signal",
     "ROOT CAUSE: No automated signal engine existed\n"
     "FIX: 6-factor scoring engine with smart strike selection\n"
     "RESULT: Automatic BUY CE / BUY PE / Iron Condor / NO TRADE with exact strikes"),
    ("No GEX (Gamma Exposure)",
     "ROOT CAUSE: GEX requires institutional math — not available in free tools\n"
     "FIX: Black-Scholes gamma × OI × lot size × spot → GEX in Crores\n"
     "RESULT: Gamma Wall, Flip Level, and regime identification in real-time"),
]

tbl = doc.add_table(rows=0, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header
hrow = tbl.add_row()
for i, h in enumerate(["Problem", "Root Cause & Solution"]):
    c = hrow.cells[i]
    set_cell_bg(c, "0D1B2A")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_WHITE

for i, (prob, sol) in enumerate(problems):
    row = tbl.add_row()
    bg = "EAF0F8" if i % 2 == 0 else "F8FBFF"
    c0, c1 = row.cells[0], row.cells[1]

    set_cell_bg(c0, bg)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    r0 = p0.add_run(f"❌  {prob}")
    r0.font.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = C_RED

    set_cell_bg(c1, "FFFFFF" if i % 2 == 0 else "FAFCFF")
    first_line = True
    for line in sol.split("\n"):
        if first_line:
            p1 = c1.paragraphs[0]
            first_line = False
        else:
            p1 = c1.add_paragraph()
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(2)
        color = C_GREEN if line.startswith("RESULT") else C_BLUE if line.startswith("FIX") else C_MID_GRAY
        bold  = line.startswith("FIX") or line.startswith("RESULT")
        r1 = p1.add_run(line)
        r1.font.size  = Pt(9)
        r1.font.bold  = bold
        r1.font.color.rgb = color

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 3 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 2 — System Architecture & File Structure")

arch_tbl = doc.add_table(rows=1, cols=2)
arch_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Left: file tree
c_left = arch_tbl.rows[0].cells[0]
set_cell_bg(c_left, "0D1B2A")
p = c_left.paragraphs[0]
p.paragraph_format.space_before = Pt(6)
r = p.add_run("File Structure")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_ACCENT

tree = [
    "nse_fo_system/",
    "  ├── web_dashboard.py      ← Main UI",
    "  ├── start_web.bat         ← One-click start",
    "  ├── get_token.py          ← Zerodha login",
    "  ├── HOW_TO_RUN.md",
    "  ├── GEX_TRADING_GUIDE.md",
    "  │",
    "  ├── config/",
    "  │   └── settings.py       ← API keys, thresholds",
    "  │",
    "  ├── core/",
    "  │   ├── kite_manager.py   ← API wrapper + cache",
    "  │   ├── pcr_tracker.py    ← PCR + OI chain",
    "  │   ├── max_pain.py       ← Max pain calc",
    "  │   ├── uoa_scanner.py    ← Unusual activity",
    "  │   ├── greeks.py         ← Black-Scholes",
    "  │   ├── risk_manager.py   ← Portfolio risk",
    "  │   └── market_utils.py   ← Expiry, lot size",
    "  │",
    "  ├── data/",
    "  │   ├── kite_token.pkl    ← Daily access token",
    "  │   ├── prev_oi.pkl       ← OI snapshot (persist)",
    "  │   └── trade_log.py      ← Trade journal",
    "  │",
    "  └── logs/",
]

for line in tree:
    if first := (line == tree[0]):
        pt = c_left.paragraphs[0]
        pt.add_run("\n")
        first = False
    else:
        pt = c_left.add_paragraph()
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after  = Pt(0)
    run = pt.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x88, 0xFF, 0xAA)

# Right: data flow
c_right = arch_tbl.rows[0].cells[1]
set_cell_bg(c_right, "F2F4F7")
p2 = c_right.paragraphs[0]
p2.paragraph_format.space_before = Pt(6)
r2 = p2.add_run("Data Flow — Every 60 Seconds")
r2.font.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = C_BLUE

flow_steps = [
    ("1", "005CA8", "Zerodha Kite API", "LTP prices, VIX, Option chain OI"),
    ("2", "006644", "Parallel Fetch (7 workers)", "Prices, OI chain, PCR, Max Pain, UOA, IV, Risk — all at once"),
    ("3", "7B2D00", "GEX Calculator", "Gamma × OI × Lot × Spot → Gamma Wall + Flip Level"),
    ("4", "005CA8", "Trade Signal Engine", "6-factor scoring → Smart strike selection"),
    ("5", "006644", "Streamlit Render", "Browser UI updates — page does NOT reload"),
]

for num, bg, title, desc in flow_steps:
    pf = c_right.add_paragraph()
    pf.paragraph_format.space_before = Pt(5)
    pf.paragraph_format.space_after  = Pt(2)
    r_num = pf.add_run(f"  [{num}]  ")
    r_num.font.bold = True
    r_num.font.size = Pt(10)
    r_num.font.color.rgb = RGBColor(int(bg[:2], 16), int(bg[2:4], 16), int(bg[4:], 16))
    r_title = pf.add_run(title)
    r_title.font.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = C_TEXT
    pd = c_right.add_paragraph()
    pd.paragraph_format.space_before = Pt(0)
    pd.paragraph_format.space_after  = Pt(3)
    rd = pd.add_run(f"      {desc}")
    rd.font.size = Pt(9)
    rd.font.color.rgb = C_MID_GRAY

doc.add_paragraph()
section_divider("  Key Technical Innovations", bg_hex="0A3A2A")

tech_rows = [
    ("30-min Instruments Cache",
     "kite.instruments('NFO') called ONCE every 30 minutes instead of every refresh.\n"
     "Impact: Eliminated the #1 bottleneck — saved 10-15 seconds per refresh cycle."),
    ("ThreadPoolExecutor (7 workers)",
     "All 7 API calls (prices, OI, PCR, max pain, UOA, IV, risk) run in PARALLEL.\n"
     "Impact: Sequential 15–25s → Parallel 3–5s (5x improvement)."),
    ("prev_oi.pkl Persistence",
     "OI snapshot saved to disk after every fetch using Python pickle.\n"
     "Impact: OI Change (CE CHG / PE CHG) shows real values even after restart."),
    ("@st.fragment(run_every=60)",
     "Streamlit native fragment — only the data section re-renders every 60s.\n"
     "Impact: Page never reloads. Sidebar, filters, connection state preserved."),
    ("Black-Scholes Gamma for GEX",
     "Uses existing greeks.py calc_greeks() — no new API calls needed for GEX.\n"
     "Impact: GEX calculated from already-fetched OI data — zero extra cost."),
]
two_col_table(tech_rows, header=["Innovation", "How It Works & Impact"])
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 4 — GEX
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 3 — GEX (Gamma Exposure) — The Unique Feature")

body_para(
    "GEX is an institutional-grade signal used by hedge funds and prop desks. "
    "It is NOT available in Sensibull, Opstra, or any other retail tool in India. "
    "This dashboard calculates it in real-time using live OI data.",
    size=11, color=C_TEXT
)

# GEX explanation table
gex_tbl = doc.add_table(rows=1, cols=3)
gex_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

gex_concepts = [
    ("005CA8", "📐  Formula",
     "GEX (Cr) = Gamma × OI × Lot Size × Spot Price ÷ 1,00,00,000\n\n"
     "CE options → Positive GEX (stabilizing)\n"
     "PE options → Negative GEX (destabilizing)\n\n"
     "Net GEX = Sum of all CE GEX − Sum of all PE GEX"),
    ("006644", "📦  Positive GEX",
     "Net GEX > +0.5 Cr → RANGE BOUND\n\n"
     "Market makers are NET sellers of options.\n"
     "Their hedging REDUCES market moves.\n\n"
     "Price goes up → MM sells futures → price comes back\n"
     "Price goes down → MM buys futures → price comes back\n\n"
     "Best strategy: SELL premium (Iron Condor)"),
    ("880000", "🌊  Negative GEX",
     "Net GEX < −0.5 Cr → VOLATILE / TRENDING\n\n"
     "Market makers are NET buyers of options.\n"
     "Their hedging AMPLIFIES market moves.\n\n"
     "Price goes up → MM also buys → goes further up\n"
     "Price goes down → MM also sells → goes further down\n\n"
     "Best strategy: BUY options (directional)"),
]

for i, (bg, title, text) in enumerate(gex_concepts):
    c = gex_tbl.rows[0].cells[i]
    set_cell_bg(c, bg)
    ph = c.paragraphs[0]
    ph.paragraph_format.space_before = Pt(8)
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run(title)
    rh.font.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = C_WHITE
    pt = c.add_paragraph()
    pt.paragraph_format.space_after = Pt(8)
    rt = pt.add_run(f"\n{text}")
    rt.font.size = Pt(9.5); rt.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)

doc.add_paragraph()

# Gamma Wall & Flip Level
section_divider("  Gamma Wall & Flip Level", bg_hex="0A3A2A")

gw_fl = doc.add_table(rows=1, cols=2)
gw_fl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (bg, title, explanation) in enumerate([
    ("1A3A6A",
     "⚡  Gamma Wall",
     [
         "Definition: The strike with the HIGHEST absolute net GEX.",
         "",
         "Why it matters: Market makers have the most exposure here.",
         "Their hedging activity creates the strongest buying/selling pressure.",
         "",
         "Behavior: Price is magnetically attracted to this strike.",
         "It acts as BOTH strongest support AND strongest resistance.",
         "",
         "Trading use:",
         "  • Price 50–100 pts below wall → SELL CE at wall",
         "  • Price 50–100 pts above wall → SELL PE at wall",
         "  • Wall breakout + volume → Buy in breakout direction",
     ]),
    ("1A4A2A",
     "🔁  Flip Level",
     [
         "Definition: The strike where cumulative GEX crosses ZERO.",
         "",
         "Why it matters: It's the boundary between 'safe' and 'danger'.",
         "",
         "ABOVE Flip Level (Safe Zone):",
         "  → Market is stable",
         "  → Option selling strategies work well",
         "  → Iron Condor, Short Strangle — safe",
         "",
         "BELOW Flip Level (Danger Zone):",
         "  → Market can become volatile",
         "  → Avoid naked option selling",
         "  → Tighten stop losses immediately",
         "  → Option buying preferred",
     ]),
]):
    c = gw_fl.rows[0].cells[i]
    set_cell_bg(c, bg)
    ph = c.paragraphs[0]
    ph.paragraph_format.space_before = Pt(8)
    rh = ph.add_run(title)
    rh.font.bold = True; rh.font.size = Pt(12); rh.font.color.rgb = C_WHITE
    for line in explanation:
        pl = c.add_paragraph()
        pl.paragraph_format.space_before = Pt(1)
        pl.paragraph_format.space_after  = Pt(1)
        rl = pl.add_run(line)
        rl.font.size = Pt(9.5)
        rl.font.color.rgb = RGBColor(0xCC, 0xEE, 0xFF) if line.startswith("  ") else RGBColor(0xDD, 0xEE, 0xFF)
        if line.startswith("ABOVE") or line.startswith("BELOW") or line.startswith("Definition") or line.startswith("Why") or line.startswith("Behavior") or line.startswith("Trading"):
            rl.font.bold = True

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 5 — TRADE SIGNAL ENGINE
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 4 — Trade Signal Engine — 6-Factor Analysis")

body_para(
    "The signal engine analyzes 6 independent factors, assigns a score to each, "
    "and combines them into a single Confidence Score (0–100%). "
    "A minimum score of 30 is required before any trade signal is generated.",
    size=11, color=C_TEXT
)

factors_data = [
    ("1", "📉 PCR — Put-Call Ratio",
     "PCR > 1.2 (Bullish) = +20 pts\nPCR < 0.8 (Bearish) = -20 pts\nTrend ▲ = +15 pts, Trend ▼ = -15 pts",
     "Measures sentiment. Rising PCR = Bears protecting = Bulls in control"),
    ("2", "🏗️ OI Build Signal",
     "CE OI Change > 500 (Fresh Long) = +25 pts\nPE OI Change > 500 (Fresh Short) = -25 pts\nLong Unwind = -10 pts, Short Cover = +10 pts",
     "Tracks NEW money entering the market. Most reliable real-time signal"),
    ("3", "😨 VIX — Fear Index",
     "VIX < 15 (Low fear) = +10 pts, switch to Buy mode\nVIX > 20 (High fear) = -10 pts, switch to Sell mode",
     "Low VIX = Cheap options = Good time to buy. High VIX = Expensive = Sell"),
    ("4", "📊 IV Rank",
     "IV Rank < 30% (Cheap) = +10 pts, Buy confirmed\nIV Rank > 60% (Expensive) = switch to Sell mode",
     "Compares current IV to historical range. Cheap IV = Buy options is smart"),
    ("5", "⚡ GEX Regime",
     "VOLATILE (Negative GEX) = ±10 pts (amplifies direction)\nRANGE BOUND (Positive GEX) = -10 pts, switch to Sell mode",
     "Tells whether market maker hedging will amplify or dampen moves"),
    ("6", "💊 Max Pain",
     "Price near Max Pain (< 2 strikes) = Balanced\nPrice above Max Pain = may pull down\nPrice below Max Pain = may pull up",
     "Options expiry gravitational pull. Strongest on expiry week"),
]

for num, name, scoring, logic in factors_data:
    ft = doc.add_table(rows=1, cols=3)
    ft.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_bg(ft.rows[0].cells[0], "0D1B2A")
    set_cell_bg(ft.rows[0].cells[1], "EAF0F8")
    set_cell_bg(ft.rows[0].cells[2], "F8FBFF")

    c0 = ft.rows[0].cells[0]
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    r0 = p0.add_run(f"\n  {name}")
    r0.font.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = C_WHITE

    c1 = ft.rows[0].cells[1]
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    r1h = p1.add_run("Scoring:\n")
    r1h.font.bold = True; r1h.font.size = Pt(9); r1h.font.color.rgb = C_BLUE
    r1 = p1.add_run(scoring)
    r1.font.size = Pt(9); r1.font.color.rgb = C_TEXT

    c2 = ft.rows[0].cells[2]
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    r2h = p2.add_run("Why it matters:\n")
    r2h.font.bold = True; r2h.font.size = Pt(9); r2h.font.color.rgb = C_GREEN
    r2 = p2.add_run(logic)
    r2.font.size = Pt(9); r2.font.color.rgb = C_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

doc.add_paragraph()
section_divider("  Smart Strike Selection Logic", bg_hex="0A3A2A")

strike_tbl = doc.add_table(rows=0, cols=3)
strike_tbl.style = 'Table Grid'
# header
hrow = strike_tbl.add_row()
for i, h in enumerate(["Condition", "Strike Chosen", "Reasoning"]):
    c = hrow.cells[i]
    set_cell_bg(c, "005CA8")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = C_WHITE

strike_rows = [
    ("Score ≥ 55  AND  IV Rank < 30%  AND  VIX < 18", "1 OTM Strike", "Strong signal + cheap IV = OTM gives more leverage. Risk/reward better."),
    ("VIX > 20 (High fear)", "ATM Strike", "High volatility = ATM safer. OTM can lose value even if direction correct."),
    ("Price within 2 strikes of Gamma Wall", "Gamma Wall Strike", "Strongest price magnet. Institutional hedging concentrated here."),
    ("All other cases", "ATM Strike", "Safest default. Highest liquidity, most responsive to market moves."),
    ("Sell Mode (IV Rank > 55%)", "Top CE OI Strike (CE)\nTop PE OI Strike (PE)", "Sell at resistance/support where most open interest is concentrated."),
]

for i, (cond, strike, reason) in enumerate(strike_rows):
    row = strike_tbl.add_row()
    bg = "F2F4F7" if i % 2 == 0 else "FFFFFF"
    for j, (cell_data, bg_c) in enumerate([(cond, "EAF0F8"), (strike, "FFFDE7"), (reason, bg)]):
        c = row.cells[j]
        set_cell_bg(c, bg_c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(cell_data)
        run.font.size = Pt(9)
        run.font.bold = (j == 1)
        colors = [C_BLUE, C_GOLD, C_TEXT]
        run.font.color.rgb = colors[j]

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 6 — DASHBOARD PANELS
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 5 — Dashboard Panels Overview")

panels = [
    ("🌐 Market Overview",
     "NIFTY / BANKNIFTY / FINNIFTY live prices, India VIX, market status (Open/Pre-open/Closed). "
     "Color-coded: green for up, red for down."),
    ("🎯 Trade Signal (Top priority panel)",
     "Automatic BUY CE / BUY PE / SELL Iron Condor / NO TRADE with:\n"
     "• Exact strike price with reason WHY that strike was chosen\n"
     "• Entry price, Target (+42%), Stop Loss (-28%), Number of lots\n"
     "• Max profit in ₹ and max loss in ₹\n"
     "• 6-factor checklist (✅ ❌ ⚠️) — every factor visible\n"
     "• Confidence score (0–100%) with visual bar"),
    ("📊 OI Chain",
     "ATM ±10 strikes with CE OI, CE CHG, CE LTP | PCR | PE LTP, PE CHG, PE OI.\n"
     "• ATM row highlighted in green (white text)\n"
     "• OI in Lakh format (e.g. 197.5L instead of 19,750,000)\n"
     "• CE CHG / PE CHG shows change from PREVIOUS refresh (persisted to disk)"),
    ("🔍 Unusual Options Activity (UOA)",
     "Scans NIFTY & BANKNIFTY for volume 2x+ above average.\n"
     "🔥 = 5x+ (Fire signal), 📈 = 2x+ (Unusual). Shows symbol, strike, type, volume ratio."),
    ("📉 PCR Readings",
     "PCR for NIFTY and BANKNIFTY with zone (Extreme Bull / Bullish / Neutral / Bearish / Extreme Bear), "
     "signal (STRONG BUY / BUY / SIDEWAYS / SELL / STRONG SELL), and recommended strategy."),
    ("🎯 IV Rank · Greeks · Skew",
     "ATM IV %, IV Rank %, Delta, Gamma, Theta (₹/day), Vega, and IV Skew (PE-CE IV difference). "
     "Skew > 0 = Puts more expensive = market expects downside."),
    ("⚡ GEX — Gamma Exposure",
     "Regime card (Range Bound / Volatile / Neutral), Gamma Wall strike, Flip Level strike, "
     "per-strike GEX table (CE GEX, PE GEX, Net GEX), and 'How to use' expander."),
    ("🏗️ OI Buildup Analysis",
     "Summary of FRESH LONG / FRESH SHORT / LONG UNWIND / SHORT COVER at each strike. "
     "Identifies smart money entry/exit points."),
    ("🛡️ Portfolio Risk",
     "Live P&L across open positions, delta exposure, margin utilization, daily loss limit."),
]

for name, desc in panels:
    pt = doc.add_table(rows=1, cols=1)
    cell = pt.rows[0].cells[0]
    set_cell_bg(cell, "EAF4FF")
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_before = Pt(4)
    rh = ph.add_run(f"  {name}")
    rh.font.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = C_BLUE
    pd_ = cell.add_paragraph()
    pd_.paragraph_format.space_before = Pt(2)
    pd_.paragraph_format.space_after  = Pt(5)
    rd = pd_.add_run(f"  {desc}")
    rd.font.size = Pt(9.5); rd.font.color.rgb = C_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 7 — TRADING STRATEGIES
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 6 — Trading Strategies Generated by the System")

strat_tbl = doc.add_table(rows=0, cols=4)
strat_tbl.style = 'Table Grid'
# Header
hrow = strat_tbl.add_row()
for i, h in enumerate(["Signal", "Setup Required", "Strike / Action", "Risk Parameters"]):
    c = hrow.cells[i]
    set_cell_bg(c, "0D1B2A")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_WHITE

strategies = [
    ("🟢 BUY CE\n(Bullish)",
     "• GEX Negative (Volatile)\n• PCR > 1.2 and rising\n• OI Build = Fresh Long\n• Price > Flip Level",
     "ATM CE (default)\nor 1 OTM if Score ≥ 55 + IV < 30%\nor Gamma Wall strike",
     "Entry: Market\nTarget: +42%\nSL: -28%\nLots: ₹2000 max risk\nTimeframe: Intraday"),
    ("🔴 BUY PE\n(Bearish)",
     "• GEX Negative (Volatile)\n• PCR < 0.8 and falling\n• OI Build = Fresh Short\n• Price < Flip Level",
     "ATM PE (default)\nor 1 OTM if Score ≥ 55 + IV < 30%\nor Gamma Wall strike",
     "Entry: Market\nTarget: +42%\nSL: -28%\nLots: ₹2000 max risk\nTimeframe: Intraday"),
    ("💰 SELL Iron Condor\n(Range Bound)",
     "• GEX Positive (Range Bound)\n• IV Rank > 55%\n• PCR Neutral (0.8–1.2)\n• VIX > 15",
     "SELL CE: Top CE OI strike or above Gamma Wall\nSELL PE: Top PE OI strike\n(Natural resistance/support)",
     "Max Profit: Total premium × lot\nSL: Either side 2x premium\nTimeframe: Weekly\nRequires: 4-leg setup"),
    ("⛔ NO TRADE\n(Wait)",
     "• Score < 30\n• Mixed signals\n• Setup not confirmed",
     "No action taken.\nFactor checklist shown\nso trader can see what's\nmissing for a signal.",
     "Capital preserved.\nWait for clearer setup.\nRe-check next 60s\nauto-refresh."),
]

for i, (sig, setup, action, risk) in enumerate(strategies):
    row = strat_tbl.add_row()
    colors_bg = ["0A2A0A", "2A1000", "0A1A2A", "1E2130"]
    colors_text = [C_GREEN, C_RED, C_ACCENT, C_MID_GRAY]
    set_cell_bg(row.cells[0], colors_bg[i])
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(sig)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = colors_text[i]

    for j, (cell_data, bg_c) in enumerate([(setup, "EAF0F8"), (action, "FFFDE7"), (risk, "F8FBFF")]):
        c = row.cells[j + 1]
        set_cell_bg(c, bg_c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(cell_data)
        r.font.size = Pt(9); r.font.color.rgb = C_TEXT

doc.add_paragraph()
section_divider("  Risk Management Rules", bg_hex="6B0000")

risk_rows = [
    ("Max loss per trade",      "₹2,000 (2% of ₹1L capital). Lot size calculated automatically to respect this."),
    ("Stop Loss — Options Buy", "−28% of premium. If ₹185 entry → SL at ₹133"),
    ("Target — Options Buy",    "+42% of premium. If ₹185 entry → Target ₹263"),
    ("Stop Loss — Iron Condor", "Exit if either leg premium becomes 2× collected premium"),
    ("Regime change exit",      "If GEX regime flips after entry — exit immediately regardless of P&L"),
    ("Max open positions",      "3 (configured in settings.py)"),
    ("Daily max loss",          "₹10,000 (configured in settings.py)"),
    ("Break-even win rate",     "With 1.5:1 R:R ratio, need only 40% winning trades to be profitable"),
]
two_col_table(risk_rows, header=["Rule", "Detail"])
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 8 — COMPARISON + KNOWN ISSUES + HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 7 — Competitive Comparison")

comp_tbl = doc.add_table(rows=0, cols=5)
comp_tbl.style = 'Table Grid'
hrow = comp_tbl.add_row()
headers = ["Feature", "This Dashboard", "Sensibull", "Opstra", "Broker App"]
bg_h    = ["0D1B2A",  "005CA8",         "555555",    "555555", "555555"]
for i, (h, bg) in enumerate(zip(headers, bg_h)):
    c = hrow.cells[i]
    set_cell_bg(c, bg)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_WHITE

comp_rows = [
    ("GEX (Gamma Exposure)",       "✅ Real-time",   "❌ No",  "❌ No",  "❌ No"),
    ("Gamma Wall",                 "✅ Real-time",   "❌ No",  "❌ No",  "❌ No"),
    ("Flip Level",                 "✅ Real-time",   "❌ No",  "❌ No",  "❌ No"),
    ("Auto BUY/SELL Signal",       "✅ With strike", "⚠️ Paid","⚠️ Paid","❌ No"),
    ("6-Factor Confidence Score",  "✅ Yes",         "❌ No",  "❌ No",  "❌ No"),
    ("Smart Strike Selection",     "✅ Yes",         "❌ No",  "❌ No",  "❌ No"),
    ("OI Change Persistence",      "✅ Disk-saved",  "⚠️ Session","⚠️ Session","❌ No"),
    ("Cost",                       "✅ Free",        "❌ ₹999+/mo","❌ ₹499+/mo","⚠️ Broker fees"),
    ("Customizable",               "✅ Full control","❌ No",  "❌ No",  "❌ No"),
    ("Data Source",                "✅ Live Zerodha","⚠️ NSE delayed","⚠️ NSE delayed","✅ Live"),
]

for i, row_data in enumerate(comp_rows):
    row = comp_tbl.add_row()
    bg = "F2F4F7" if i % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], "EAF0F8")
    p = row.cells[0].paragraphs[0]
    r = p.add_run(row_data[0])
    r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = C_BLUE
    for j in range(1, 5):
        set_cell_bg(row.cells[j], "E8F5E9" if "✅" in row_data[j] else ("FFF8E1" if "⚠️" in row_data[j] else "FFEBEE"))
        p = row.cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(row_data[j])
        r.font.size = Pt(9.5)
        r.font.color.rgb = C_GREEN if "✅" in row_data[j] else (C_GOLD if "⚠️" in row_data[j] else C_RED)

doc.add_paragraph()
section_divider("  SECTION 8 — Known Issues & Planned Fixes", bg_hex="6B0000")

issues = [
    ("CRITICAL — IV Rank always ~33%",
     "CURRENT: Formula calculates rank relative to current IV itself → always 33.3%\n"
     "FIX NEEDED: Store daily ATM IV in data/iv_history.json, calculate rank from 52-week range\n"
     "IMPACT: Sell mode rarely triggers from IV alone; OTM selection condition unreliable"),
    ("MEDIUM — OI Build threshold absolute (500)",
     "CURRENT: Fixed 500 contracts regardless of market activity\n"
     "FIX NEEDED: Use relative threshold = 5% of total OI at that strike\n"
     "IMPACT: Expiry week generates too many false FL/FS signals"),
    ("LOW — No time-of-day filter",
     "CURRENT: BUY signals generated at 3:20 PM (30 min before close)\n"
     "FIX NEEDED: Disable buy signals after 2:45 PM, show warning\n"
     "IMPACT: Overnight theta decay on options bought in last 30 min"),
    ("LOW — Score thresholds not backtested",
     "CURRENT: 30/55 thresholds chosen by logic, not historical data\n"
     "FIX NEEDED: Backtest on 6-month historical data to find optimal thresholds\n"
     "IMPACT: May generate signals that have < 40% win rate at current thresholds"),
]

for title, desc in issues:
    it = doc.add_table(rows=1, cols=1)
    cell = it.rows[0].cells[0]
    level_bg = "3A0000" if "CRITICAL" in title else ("2A1A00" if "MEDIUM" in title else "1A2A1A")
    set_cell_bg(cell, level_bg)
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_before = Pt(4)
    rh = ph.add_run(f"  ⚠️  {title}")
    rh.font.bold = True; rh.font.size = Pt(10)
    rh.font.color.rgb = C_RED if "CRITICAL" in title else (C_GOLD if "MEDIUM" in title else C_GREEN)
    pd_ = cell.add_paragraph()
    pd_.paragraph_format.space_after = Pt(4)
    rd = pd_.add_run(f"  {desc.replace(chr(10), chr(10) + '  ')}")
    rd.font.size = Pt(9); rd.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 9 — HOW TO RUN + QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

section_divider("  SECTION 9 — How to Run the Dashboard")

run_tbl = doc.add_table(rows=1, cols=2)

# Left: daily steps
cl = run_tbl.rows[0].cells[0]
set_cell_bg(cl, "0D1B2A")
ph = cl.paragraphs[0]
ph.paragraph_format.space_before = Pt(8)
rh = ph.add_run("  Daily Routine (Every Market Day)")
rh.font.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = C_ACCENT

steps = [
    ("FIRST TIME ONLY", "pip install streamlit kiteconnect scipy numpy pandas"),
    ("Step 1", "Double-click start_web.bat"),
    ("Step 2", "If token prompt: type Y, browser opens for Zerodha login"),
    ("Step 3", "Copy URL after login, paste in CMD"),
    ("Step 4", "Dashboard opens at http://localhost:8501"),
    ("Step 5", "Select symbol (NIFTY / BANKNIFTY) from sidebar"),
    ("Step 6", "Check Trade Signal panel — act only if ✅ majority"),
    ("Stop", "Press Ctrl+C in CMD window"),
]

for step, action in steps:
    ps = cl.add_paragraph()
    ps.paragraph_format.space_before = Pt(4)
    ps.paragraph_format.space_after  = Pt(2)
    rs = ps.add_run(f"  {step}: ")
    rs.font.bold = True; rs.font.size = Pt(9.5); rs.font.color.rgb = C_GOLD
    ra = ps.add_run(action)
    ra.font.size = Pt(9.5); ra.font.color.rgb = RGBColor(0xCC, 0xEE, 0xFF)

# Right: quick reference card
cr = run_tbl.rows[0].cells[1]
set_cell_bg(cr, "0A2A0A")
ph2 = cr.paragraphs[0]
ph2.paragraph_format.space_before = Pt(8)
rh2 = ph2.add_run("  Quick Reference Card")
rh2.font.bold = True; rh2.font.size = Pt(11); rh2.font.color.rgb = C_GREEN

qr_items = [
    ("GEX POSITIVE", "= Range Bound = SELL options"),
    ("GEX NEGATIVE", "= Volatile = BUY options"),
    ("GAMMA WALL",   "= Strongest support/resistance"),
    ("FLIP LEVEL",   "= Safe zone boundary"),
    ("Spot > Flip",  "= Safe to sell premium"),
    ("Spot < Flip",  "= Buy options / tight SL"),
    ("PCR > 1.2 ▲",  "= Bullish sentiment"),
    ("PCR < 0.8 ▼",  "= Bearish sentiment"),
    ("VIX < 15",     "= Options cheap, ok to buy"),
    ("VIX > 20",     "= Options expensive, sell them"),
    ("IV Rank < 30%","= Cheap IV, buy options"),
    ("IV Rank > 60%","= Expensive IV, sell options"),
    ("Score ≥ 30",   "= Trade signal generated"),
    ("Score < 30",   "= NO TRADE, wait"),
]

for key, val in qr_items:
    pq = cr.add_paragraph()
    pq.paragraph_format.space_before = Pt(3)
    pq.paragraph_format.space_after  = Pt(1)
    rk = pq.add_run(f"  {key:15s}")
    rk.font.bold = True; rk.font.name = "Courier New"; rk.font.size = Pt(9)
    rk.font.color.rgb = C_GOLD
    rv = pq.add_run(val)
    rv.font.size = Pt(9); rv.font.color.rgb = RGBColor(0xCC, 0xFF, 0xCC)

doc.add_paragraph()

# Final note
info_box([
    "📘  Documentation Files Included:",
    "     HOW_TO_RUN.md  —  Step-by-step setup guide, common errors, daily checklist",
    "     GEX_TRADING_GUIDE.md  —  11-part guide: GEX concepts, 6 strategies, real examples, decision tree",
    "",
    "⚠️  Disclaimer: This dashboard is an analysis tool. All trading decisions are the responsibility of the trader.",
    "     No tool can guarantee profit. Always use proper stop losses. Past signal accuracy does not guarantee future results.",
], bg_hex="1A1A2E")


# ── Save ───────────────────────────────────────────────────────────────────────
output = r"D:\HDFC\nse_fo_system\NSE_FO_Dashboard_Presentation.docx"
doc.save(output)
print(f"Document saved: {output}")
