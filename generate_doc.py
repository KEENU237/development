from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # headings
C_MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings
C_ACCENT     = RGBColor(0x00, 0x70, 0xC0)   # labels
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
C_TABLE_HDR  = RGBColor(0x1F, 0x49, 0x7D)
C_BODY       = RGBColor(0x26, 0x26, 0x26)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   kwargs.get('val',   'single'))
        border.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'CCCCCC'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(24)
    run.font.color.rgb = C_DARK_BLUE
    p.space_after = Pt(4)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = C_MID_BLUE
    run.italic = True
    p.space_after = Pt(2)

def add_label(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_ACCENT
    run.bold = True
    p.space_after = Pt(14)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK_BLUE
    p.space_before = Pt(18)
    p.space_after  = Pt(6)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_MID_BLUE
    p.space_before = Pt(12)
    p.space_after  = Pt(4)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_ACCENT
    p.space_before = Pt(8)
    p.space_after  = Pt(3)

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_BODY
    p.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_BODY
    p.space_after = Pt(3)

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EBF3FB')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1F497D')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = 'F2F2F2' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run  = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = C_BODY

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().space_after = Pt(6)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.space_before = Pt(4)
    p.space_after  = Pt(10)

def add_callout(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EBF3FB')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = C_DARK_BLUE
    run.italic = True

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── Cover ─────────────────────────────────────────────────────────────────────
doc.add_paragraph().space_after = Pt(30)
add_title(doc, "File-Based Data Ingestion Architecture")
add_subtitle(doc, "From Cloud Storage to Your Data Platform")
add_label(doc, "Client Presentation Document  |  Final Version  |  April 2026")
add_divider(doc)

# ── Executive Summary ─────────────────────────────────────────────────────────
add_h1(doc, "EXECUTIVE SUMMARY")
add_callout(doc,
    "Your business receives data files every day — from internal systems, partners, banks, and "
    "third-party vendors. These files sit in cloud storage (Amazon S3) in raw, unprocessed form.")
add_body(doc,
    "The challenge: Raw files cannot go directly into Snowflake. They are the wrong format, "
    "wrong size, and often contain dirty data.")
add_body(doc,
    "Our solution: An automated processing layer that sits between S3 and Snowflake — cleans "
    "every file, converts it to the right format, splits it to the right size, and loads it "
    "automatically. Zero manual effort after setup.")
add_divider(doc)

# ── Complete Picture ──────────────────────────────────────────────────────────
add_h1(doc, "THE COMPLETE PICTURE")
add_code_block(doc,
"                    YOUR DATA JOURNEY\n\n"
"  ┌─────────────┐     ┌──────────────────┐     ┌─────────────┐\n"
"  │             │     │                  │     │             │\n"
"  │  S3 BUCKET  │────▶│  PROCESSING      │────▶│  SNOWFLAKE  │\n"
"  │             │     │  LAYER           │     │             │\n"
"  │  Raw Files  │     │  AWS Glue        │     │  Clean Data │\n"
"  │  Any Format │     │  Python Shell    │     │  Ready for  │\n"
"  │  Any Size   │     │                  │     │  Reports    │\n"
"  └─────────────┘     └──────────────────┘     └─────────────┘\n\n"
"      ARRIVE                PROCESS                  LOAD\n"
"   Automatically          Automatically           Automatically")
add_body(doc, "Every step is automated. Your team receives clean, accurate data — on time, every time.")
add_divider(doc)

# ── Section 1 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 1 — FILES IN S3")
add_h2(doc, "What Is Amazon S3?")
add_body(doc,
    "Amazon S3 is a secure cloud storage service — like a shared drive on the internet, "
    "but enterprise-grade, highly secure, and infinitely scalable. Your files land here "
    "automatically from various business sources.")

add_h2(doc, "Types of Files We Handle")
add_table(doc,
    ["File Type", "Where It Typically Comes From"],
    [
        ["CSV",     "ERP systems, bank statements, POS exports, legacy tools"],
        ["JSON",    "Web applications, mobile apps, REST APIs"],
        ["Excel",   "Finance teams, manual business uploads, partner reports"],
        ["XML",     "Government portals, insurance systems, legacy enterprise software"],
        ["Parquet", "Data warehouse exports, analytics platform outputs"],
    ],
    col_widths=[1.4, 4.5])

add_h2(doc, "The Problem With Raw Files")
add_body(doc,
    "Raw files arriving in S3 have three common problems that prevent direct loading into Snowflake:")

add_h3(doc, "Problem 1 — Wrong Format")
add_body(doc,
    "Snowflake works best with Parquet format. Most business systems export CSV or Excel. "
    "Loading raw CSV into Snowflake is like printing a document to read it on screen — "
    "it works, but it is slow and wasteful.")

add_h3(doc, "Problem 2 — Wrong Size")
add_table(doc,
    ["Situation", "What Happens"],
    [
        ["One giant 10 GB file",           "Snowflake loads it as a single thread — very slow"],
        ["50,000 files of 1 MB each",      "Snowflake is overwhelmed managing file metadata — very slow"],
        ["100–250 MB files  ✅ Optimal",   "Snowflake loads all files simultaneously in parallel — fast"],
    ],
    col_widths=[2.5, 4.0])

add_h3(doc, "Problem 3 — Dirty Data")
add_body(doc, "Raw files frequently contain issues that break downstream reports:")
for item in [
    "Empty rows and duplicate records",
    "Inconsistent date formats — some DD/MM/YYYY, some MM-DD-YYYY",
    "Missing values in critical columns",
    "Wrong character encoding — special characters appear as symbols",
    "Column names that do not match the Snowflake table structure",
]:
    add_bullet(doc, item)
add_body(doc,
    "These problems must be fixed before data enters Snowflake. "
    "Our processing layer handles all of this automatically.")
add_divider(doc)

# ── Section 2 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 2 — THE PROCESSING LAYER")
add_h2(doc, "AWS Glue Python Shell")
add_h3(doc, "What Is It?")
add_callout(doc,
    "AWS Glue Python Shell is standard Python code running on fully managed Amazon Web Services "
    "infrastructure. There is no server to buy, configure, or maintain. Your engineering team "
    "writes regular Python — the same language they already know. AWS handles everything else — "
    "servers, scaling, patching, and availability.")
add_body(doc,
    "Most importantly: It costs nothing when it is not running. "
    "You pay only for the minutes it is actively processing your files.")

add_h2(doc, "The Four Jobs It Does")

add_h3(doc, "Job 1 — READ")
add_body(doc,
    "Opens and reads any file format from S3 — CSV, JSON, Excel, XML, or Parquet. "
    "No manual conversion needed before this step. The system accepts files exactly as they arrive.")

add_h3(doc, "Job 2 — CLEAN")
add_body(doc, "Applies a set of data quality rules to every file:")
add_table(doc,
    ["Cleaning Action", "What It Fixes"],
    [
        ["Remove duplicate rows",       "Prevents double-counting in reports"],
        ["Remove empty rows",           "Eliminates null records that break aggregations"],
        ["Standardize date formats",    "All dates converted to YYYY-MM-DD"],
        ["Fix character encoding",      "Special characters displayed correctly"],
        ["Rename columns",              "Matches exactly to Snowflake table structure"],
        ["Handle missing values",       "Replaces nulls with agreed default values"],
        ["Validate data types",         "Ensures numbers are numbers, dates are dates"],
    ],
    col_widths=[2.5, 4.0])

add_h3(doc, "Job 3 — CONVERT")
add_body(doc, "Converts the cleaned file into Parquet format — the optimal format for Snowflake.")
add_table(doc,
    ["", "CSV", "Parquet"],
    [
        ["How data is stored",        "Row by row",            "Column by column"],
        ["What Snowflake reads",       "Entire file always",    "Only columns needed"],
        ["File size",                  "Large",                 "60–80% smaller"],
        ["Load speed",                 "Baseline",              "Up to 5× faster"],
        ["Data type preservation",     "No — everything is text","Yes — preserved correctly"],
    ],
    col_widths=[2.5, 1.8, 2.3])

add_body(doc,
    "Parquet is the industry standard for cloud data platforms. Every major data platform — "
    "Snowflake, Databricks, BigQuery — recommends it.")

add_h3(doc, "Job 4 — SPLIT")
add_body(doc,
    "Breaks large converted files into 100–250 MB chunks. "
    "This is critical for Snowflake performance.")
add_code_block(doc,
"BEFORE SPLITTING:\n"
"One file of 2 GB  →  Snowflake loads in 1 thread  →  Slow\n\n"
"AFTER SPLITTING:\n"
"File split into 10 chunks of 200 MB each\n"
"→  Snowflake loads all 10 simultaneously\n"
"→  10× faster\n"
"→  Same data, same result, fraction of the time")
add_body(doc,
    "Processed files are saved back to a separate S3 folder — clean, converted, "
    "correctly sized, and ready for Snowflake.")
add_divider(doc)

# ── Section 3 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 3 — LOADING INTO SNOWFLAKE")
add_body(doc,
    "After the processing layer completes, Snowflake loads the prepared files. "
    "We use one of two loading methods depending on your business requirement.")

add_h2(doc, "Method A — Bulk Load")
add_body(doc, "Best for: Daily, hourly, or weekly batch processing")
add_code_block(doc,
"Processing completes\n"
"        ↓\n"
"Snowflake reads all processed chunks simultaneously\n"
"        ↓\n"
"All chunks loaded in parallel\n"
"        ↓\n"
"Data available in Snowflake table\n"
"        ↓\n"
"Dashboards and reports refresh automatically")
add_body(doc,
    "Performance example: A 2 GB raw CSV file — after conversion to Parquet and splitting "
    "into 10 chunks of 200 MB each — loads into Snowflake in approximately 3–5 minutes. "
    "The same file as a single raw CSV would take 25–40 minutes.")

add_h2(doc, "Method B — Snowpipe (Continuous Loading)")
add_body(doc, "Best for: Files that arrive throughout the day and must be available immediately.")
add_code_block(doc,
"New file arrives in S3\n"
"        ↓\n"
"S3 automatically notifies Snowflake\n"
"        ↓\n"
"Glue processes the file immediately\n"
"        ↓\n"
"Snowpipe loads the processed file\n"
"        ↓\n"
"Data available within 1–5 minutes of arrival\n"
"        ↓\n"
"No schedule needed — fully event-driven")

add_h2(doc, "Choosing the Right Loading Method")
add_table(doc,
    ["Requirement", "Recommended Method"],
    [
        ["Data needed once a day",                    "Bulk Load"],
        ["Data needed every hour",                    "Bulk Load"],
        ["Data needed within minutes of arrival",     "Snowpipe"],
        ["High volume — thousands of files per day",  "Bulk Load"],
        ["Unpredictable file arrival times",          "Snowpipe"],
    ],
    col_widths=[3.5, 2.5])
add_divider(doc)

# ── Section 4 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 4 — END-TO-END FLOW")
add_code_block(doc,
"┌─────────────────────────────────────────────────────────────────┐\n"
"│                        S3 — RAW ZONE                           │\n"
"│                                                                 │\n"
"│   sales_data.csv (3 GB)  ·  transactions.json  ·  report.xlsx  │\n"
"└───────────────────────────────┬─────────────────────────────────┘\n"
"                                │\n"
"                                │  File arrives → processing starts\n"
"                                ▼\n"
"┌─────────────────────────────────────────────────────────────────┐\n"
"│                  AWS GLUE PYTHON SHELL                          │\n"
"│                                                                 │\n"
"│   Step 1  READ      →  Opens any file format                    │\n"
"│   Step 2  CLEAN     →  Fixes quality issues                     │\n"
"│   Step 3  CONVERT   →  Saves as Parquet + GZIP                  │\n"
"│   Step 4  SPLIT     →  Creates 200 MB chunks                    │\n"
"└───────────────────────────────┬─────────────────────────────────┘\n"
"                                │\n"
"                                ▼\n"
"┌─────────────────────────────────────────────────────────────────┐\n"
"│                     S3 — PROCESSED ZONE                        │\n"
"│                                                                 │\n"
"│   sales_001.parquet (200 MB)                                    │\n"
"│   sales_002.parquet (200 MB)                                    │\n"
"│   sales_003.parquet (200 MB)   ← 15 chunks from original 3 GB  │\n"
"│   ...                                                           │\n"
"└───────────────────────────────┬─────────────────────────────────┘\n"
"                                │\n"
"                                │  All 15 files load simultaneously\n"
"                                ▼\n"
"┌─────────────────────────────────────────────────────────────────┐\n"
"│                         SNOWFLAKE                               │\n"
"│                                                                 │\n"
"│   ✅  Clean data in table                                       │\n"
"│   ✅  Available for reports and dashboards                      │\n"
"│   ✅  3 GB file fully loaded in under 5 minutes                 │\n"
"└─────────────────────────────────────────────────────────────────┘")
add_divider(doc)

# ── Section 5 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 5 — WHAT YOUR TEAM DOES VS WHAT IS AUTOMATED")
add_table(doc,
    ["Activity", "Automated", "Your Team"],
    [
        ["Files arrive in S3 from source systems",        "✅", ""],
        ["Processing starts automatically",               "✅", ""],
        ["Data cleaned and validated",                    "✅", ""],
        ["Files converted to Parquet",                    "✅", ""],
        ["Files split to optimal size",                   "✅", ""],
        ["Data loaded into Snowflake",                    "✅", ""],
        ["Dashboards refresh with new data",              "✅", ""],
        ["Alert sent if file fails processing",           "✅", ""],
        ["Monitor pipeline health dashboard",             "",   "✅"],
        ["Investigate and resolve exceptions",            "",   "✅"],
    ],
    col_widths=[4.0, 1.2, 1.2])
add_body(doc,
    "Your team's responsibility shifts from manually moving data to managing exceptions "
    "— which are rare once the pipeline is stable.")
add_divider(doc)

# ── Section 6 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 6 — BUSINESS BENEFITS")
add_table(doc,
    ["Benefit", "Impact"],
    [
        ["Fully automated",          "Zero manual effort for daily data movement"],
        ["5× faster loading",        "Parquet + parallel loading vs raw CSV"],
        ["70% less storage",         "Parquet compression reduces Snowflake storage cost"],
        ["Clean data guaranteed",    "Quality rules applied to every file before it reaches reports"],
        ["Scales automatically",     "10 files or 10,000 files — same architecture handles both"],
        ["Pay only when processing", "No idle server cost — Glue runs only when files need processing"],
        ["Any file format accepted", "CSV, JSON, Excel, XML — no source system changes required"],
        ["Instant failure alerts",   "Team notified immediately — no discovering problems hours later"],
    ],
    col_widths=[2.5, 4.0])
add_divider(doc)

# ── Section 7 ─────────────────────────────────────────────────────────────────
add_h1(doc, "SECTION 7 — INDICATIVE COST")
add_table(doc,
    ["Component", "Pricing Model", "Idle Cost", "Monthly Estimate"],
    [
        ["AWS Glue Python Shell", "Per minute of execution",    "Zero",    "₹3,000 – ₹20,000"],
        ["Amazon S3 Storage",     "Per GB stored",              "Minimal", "₹1,000 – ₹5,000"],
        ["Snowflake Loading",     "Included in warehouse cost", "Zero",    "Part of Snowflake plan"],
    ],
    col_widths=[2.0, 2.2, 1.1, 1.8])
add_body(doc,
    "Actual cost depends on file volume, file size, and processing frequency. "
    "A detailed cost estimate will be provided after requirements scoping.")
add_divider(doc)

# ── Summary ───────────────────────────────────────────────────────────────────
add_h1(doc, "SUMMARY")
add_table(doc,
    ["Question", "Answer"],
    [
        ["Where do files come from?",          "Amazon S3 — from any business system"],
        ["What formats are supported?",        "CSV, JSON, Excel, XML, Parquet"],
        ["What processes the files?",          "AWS Glue Python Shell — pure Python, serverless"],
        ["What format are files converted to?","Parquet with GZIP compression"],
        ["What is the optimal file size?",     "100–250 MB per chunk"],
        ["How does Snowflake load the data?",  "Bulk Load for batches, Snowpipe for real-time"],
        ["Is any manual effort required?",     "No — fully automated end to end"],
        ["What if something fails?",           "Automatic alert sent to your team instantly"],
    ],
    col_widths=[2.8, 4.0])

add_divider(doc)
add_callout(doc,
    "Files arrive in S3. AWS Glue Python Shell cleans, converts, and splits them automatically. "
    "Snowflake loads them in parallel. Your team sees accurate, up-to-date data in their reports "
    "— without touching a single file manually.")

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph().space_after = Pt(10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This document is confidential and prepared exclusively for client review.  |  "
    "A separate technical implementation guide is available for the engineering team.  |  "
    "All cost figures are indicative and subject to detailed requirements scoping.")
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Document Version: 1.0  |  Date: April 2026")
run2.font.size = Pt(8.5)
run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/development/File_Ingestion_Architecture.docx"
doc.save(out)
print(f"Saved: {out}")
